from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from docx import Document
import time
import os

print("🚀 Запускаю ChatGPT автоматизацию...")

chrome_options = Options()

# Путь к Brave
chrome_options.binary_location = "C:/Program Files/BraveSoftware/Brave-Browser/Application/brave.exe"

# НОВЫЙ профиль специально для автоматизации (не ваш основной)
profile_path = os.path.join(os.getcwd(), "brave_automation_profile")
chrome_options.add_argument(f"user-data-dir={profile_path}")

# Отключаем некоторые детекторы автоматизации
chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
chrome_options.add_experimental_option('useAutomationExtension', False)
chrome_options.add_argument("--disable-blink-features=AutomationControlled")

driver = webdriver.Chrome(
    service=Service(ChromeDriverManager().install()),
    options=chrome_options
)

# Убираем признаки webdriver
driver.execute_cdp_cmd('Network.setUserAgentOverride', {
    "userAgent": driver.execute_script("return navigator.userAgent").replace('HeadlessChrome', 'Chrome')
})

try:
    CHATGPT_URL = "https://chat.openai.com/"
    YOUR_PROMPT = "Напиши короткое стихотворение про кота"
    
    print("📱 Открываю ChatGPT...")
    driver.get(CHATGPT_URL)
    
    print("⏳ Первый запуск — нужно ВРУЧНУЮ авторизоваться в ChatGPT")
    print("⏳ У вас есть 2 МИНУТЫ чтобы залогиниться...")
    print("⏳ После авторизации скрипт продолжит работу автоматически")
    time.sleep(120)  # 2 минуты на авторизацию
    
    print("✏️ Ищу поле ввода...")
    input_box = WebDriverWait(driver, 20).until(
        EC.presence_of_element_located((By.ID, "prompt-textarea"))
    )
    
    print("📝 Вставляю текст...")
    input_box.click()
    time.sleep(1)
    input_box.send_keys(YOUR_PROMPT)
    time.sleep(1)
    
    print("🚀 Отправляю запрос...")
    input_box.send_keys(Keys.RETURN)
    
    print("⏳ Жду ответ (30-60 секунд)...")
    time.sleep(25)
    
    print("📖 Читаю ответ...")
    messages = driver.find_elements(By.CSS_SELECTOR, "[data-message-author-role='assistant']")
    
    if messages:
        last_message = messages[-1]
        response_text = last_message.text
        
        print(f"✅ Получен ответ ({len(response_text)} символов)")
        print(f"\n--- ОТВЕТ ---\n{response_text[:300]}...\n---\n")
        
        print("💾 Сохраняю в Word...")
        doc = Document()
        doc.add_heading('ChatGPT Response', 0)
        doc.add_paragraph(f"Запрос: {YOUR_PROMPT}")
        doc.add_paragraph("")
        doc.add_paragraph(response_text)
        doc.save('chatgpt_response.docx')
        
        print("🎉 Готово! Файл chatgpt_response.docx сохранен")
        print("✅ В следующий раз авторизация не понадобится!")
    else:
        print("⚠️ Не нашел ответ")
    
    print("\n⏳ Браузер останется открытым 30 секунд...")
    time.sleep(30)
    
except Exception as e:
    print(f"❌ Ошибка: {e}")
    print("⏳ Браузер останется открытым 60 секунд...")
    time.sleep(60)
    
finally:
    driver.quit()