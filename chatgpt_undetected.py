import undetected_chromedriver as uc
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from docx import Document
import time

print("🚀 Запускаю ChatGPT с защитой от детекции...")

options = uc.ChromeOptions()

# Путь к Brave (можно закомментировать если хотите использовать обычный Chrome)
options.binary_location = "C:/Program Files/BraveSoftware/Brave-Browser/Application/brave.exe"

# Создаем драйвер с защитой от детекции
driver = uc.Chrome(options=options, version_main=None)

try:
    CHATGPT_URL = "https://chat.openai.com/"
    YOUR_PROMPT = "Напиши короткое стихотворение про кота"
    
    print("📱 Открываю ChatGPT...")
    driver.get(CHATGPT_URL)
    
    print("⏳ ВАЖНО: Залогиньтесь ВРУЧНУЮ в открывшемся окне!")
    print("⏳ У вас есть 3 МИНУТЫ...")
    print("⏳ После успешного входа скрипт продолжит работу")
    time.sleep(180)  # 3 минуты на авторизацию
    
    print("✏️ Ищу поле ввода...")
    input_box = WebDriverWait(driver, 20).until(
        EC.presence_of_element_located((By.ID, "prompt-textarea"))
    )
    
    print("📝 Вставляю текст...")
    input_box.click()
    time.sleep(2)
    input_box.send_keys(YOUR_PROMPT)
    time.sleep(2)
    
    print("🚀 Отправляю запрос...")
    input_box.send_keys(Keys.RETURN)
    
    print("⏳ Жду ответ ChatGPT (может занять минуту)...")
    time.sleep(30)
    
    print("📖 Ищу ответ...")
    # Ждем пока появится хотя бы одно сообщение ассистента
    WebDriverWait(driver, 60).until(
        EC.presence_of_element_located((By.CSS_SELECTOR, "[data-message-author-role='assistant']"))
    )
    
    time.sleep(5)  # Даем время догенерировать
    
    messages = driver.find_elements(By.CSS_SELECTOR, "[data-message-author-role='assistant']")
    
    if messages:
        last_message = messages[-1]
        response_text = last_message.text
        
        print(f"✅ Получен ответ ({len(response_text)} символов)")
        print(f"\n--- ОТВЕТ ---")
        print(response_text[:500])
        print("---\n")
        
        print("💾 Сохраняю в Word...")
        doc = Document()
        doc.add_heading('ChatGPT Response', 0)
        doc.add_paragraph(f"Запрос: {YOUR_PROMPT}")
        doc.add_paragraph("")
        doc.add_paragraph(response_text)
        doc.save('chatgpt_response.docx')
        
        print("🎉 Готово! Файл chatgpt_response.docx сохранен в папке проекта")
    else:
        print("⚠️ Не удалось найти ответ")
    
    print("\n⏳ Браузер останется открытым 30 секунд...")
    time.sleep(30)
    
except Exception as e:
    print(f"❌ Ошибка: {e}")
    import traceback
    traceback.print_exc()
    print("\n⏳ Браузер останется открытым 2 минуты для проверки...")
    time.sleep(120)
    
finally:
    driver.quit()